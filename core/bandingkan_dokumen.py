"""
Tools: Bandingkan 2 versi dokumen yang sama (Word .docx dan/atau PDF,
boleh beda format) untuk melihat bagian yang berubah.

Cara kerja: teks tiap dokumen diekstrak per paragraf/baris (docx: isi
tabel ikut diekstrak, 1 baris tabel = 1 baris teks; pdf: per baris
halaman), dibandingkan baris demi baris dengan difflib, lalu baris yang
berubah dibandingkan lagi kata demi kata supaya bagian yang beda
tersorot presisi. Hasilnya ditulis sebagai laporan HTML mandiri (side-
by-side, ada ringkasan jumlah perubahan dan tombol navigasi antar-
perubahan).

Cara pakai (dari folder ini):
    python bandingkan_dokumen.py --a "versi_lama.docx" --b "versi_baru.pdf" --output "hasil.html"
"""

import argparse
import difflib
import html
import os

import docx
import fitz  # PyMuPDF
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(document):
    """
    Iterasi paragraf & tabel level body docx sesuai URUTAN ASLINYA di
    dokumen (document.paragraphs dan document.tables terpisah dan tidak
    merepresentasikan urutan gabungannya).
    """
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def extract_lines(path):
    """
    Ekstrak isi dokumen jadi daftar baris teks. docx: per paragraf, DAN
    per baris tabel (sel-selnya digabung " | ") sesuai urutan aslinya di
    dokumen - tanpa ini, konten di dalam tabel Word (mis. daftar hadir)
    akan diam-diam terlewat dari perbandingan. pdf: per baris halaman.
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        document = docx.Document(path)
        lines = []
        for block in _iter_block_items(document):
            if isinstance(block, Table):
                for row in block.rows:
                    lines.append(" | ".join(cell.text.strip() for cell in row.cells))
            else:
                lines.append(block.text)
        return lines

    if ext == ".pdf":
        lines = []
        with fitz.open(path) as doc:
            for page in doc:
                lines.extend(page.get_text().splitlines())
        return lines

    raise ValueError(f"Format file tidak didukung: {ext} (hanya .docx dan .pdf)")


def _diff_words(old_line, new_line):
    """Highlight kata yang beda di sepasang baris yang berubah (word-level diff)."""
    old_words = old_line.split(" ")
    new_words = new_line.split(" ")
    matcher = difflib.SequenceMatcher(None, old_words, new_words)

    old_parts, new_parts = [], []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        old_chunk = html.escape(" ".join(old_words[i1:i2]))
        new_chunk = html.escape(" ".join(new_words[j1:j2]))
        if tag == "equal":
            old_parts.append(old_chunk)
            new_parts.append(new_chunk)
            continue
        if old_chunk:
            old_parts.append(f'<mark class="w-del">{old_chunk}</mark>')
        if new_chunk:
            new_parts.append(f'<mark class="w-ins">{new_chunk}</mark>')

    return " ".join(old_parts), " ".join(new_parts)


def _build_rows(lines_a, lines_b):
    """Susun baris tabel side-by-side dari hasil diff, sekaligus hitung ringkasan."""
    matcher = difflib.SequenceMatcher(None, lines_a, lines_b)
    counts = {"added": 0, "removed": 0, "changed": 0}
    rows = []
    no_a = no_b = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for a_line, b_line in zip(lines_a[i1:i2], lines_b[j1:j2]):
                no_a += 1
                no_b += 1
                rows.append(("equal", no_a, html.escape(a_line), no_b, html.escape(b_line)))
        elif tag == "delete":
            for a_line in lines_a[i1:i2]:
                no_a += 1
                counts["removed"] += 1
                rows.append(("delete", no_a, html.escape(a_line), None, ""))
        elif tag == "insert":
            for b_line in lines_b[j1:j2]:
                no_b += 1
                counts["added"] += 1
                rows.append(("insert", None, "", no_b, html.escape(b_line)))
        elif tag == "replace":
            a_chunk, b_chunk = lines_a[i1:i2], lines_b[j1:j2]
            for k in range(max(len(a_chunk), len(b_chunk))):
                counts["changed"] += 1
                a_line = a_chunk[k] if k < len(a_chunk) else None
                b_line = b_chunk[k] if k < len(b_chunk) else None
                if a_line is not None:
                    no_a += 1
                if b_line is not None:
                    no_b += 1
                if a_line is not None and b_line is not None:
                    left_html, right_html = _diff_words(a_line, b_line)
                elif a_line is not None:
                    left_html, right_html = html.escape(a_line), ""
                else:
                    left_html, right_html = "", html.escape(b_line)
                rows.append(("replace", no_a if a_line is not None else None, left_html,
                              no_b if b_line is not None else None, right_html))

    return rows, counts


_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="utf-8">
<title>Perbandingan Dokumen</title>
<style>
  :root {{
    --red-bg: #ffeef0; --red-num: #ffdce0; --red-mark: #fdb8c0;
    --green-bg: #e6ffed; --green-num: #cdffd8; --green-mark: #acf2bd;
    --border: #d0d7de; --muted: #57606a;
  }}
  * {{ box-sizing: border-box; }}
  body {{
    margin: 0; font-family: "Segoe UI", system-ui, sans-serif;
    color: #1f2328; background: #f6f8fa;
  }}
  header {{
    position: sticky; top: 0; z-index: 10; background: #fff;
    border-bottom: 1px solid var(--border); padding: 14px 20px;
    display: flex; align-items: center; justify-content: space-between;
    flex-wrap: wrap; gap: 12px; box-shadow: 0 1px 3px rgba(0,0,0,.06);
  }}
  header h1 {{ font-size: 16px; margin: 0 0 4px; }}
  .doc-names {{ font-size: 13px; color: var(--muted); }}
  .summary {{ display: flex; gap: 10px; align-items: center; }}
  .badge {{ font-size: 12px; font-weight: 600; padding: 3px 10px; border-radius: 999px; }}
  .badge.added {{ background: var(--green-bg); color: #1a7f37; }}
  .badge.removed {{ background: var(--red-bg); color: #cf222e; }}
  .badge.changed {{ background: #fff8c5; color: #9a6700; }}
  .nav {{ display: flex; align-items: center; gap: 6px; }}
  .nav button {{
    border: 1px solid var(--border); background: #fff; border-radius: 6px;
    padding: 5px 10px; cursor: pointer; font-size: 13px;
  }}
  .nav button:hover {{ background: #f3f4f6; }}
  #nav-counter {{ font-size: 12px; color: var(--muted); min-width: 54px; text-align: center; }}
  table {{ border-collapse: collapse; width: 100%; table-layout: fixed; background: #fff; }}
  col.num {{ width: 46px; }}
  col.cell {{ width: 50%; }}
  td {{
    padding: 3px 10px; font-size: 13px; vertical-align: top;
    white-space: pre-wrap; word-break: break-word; border-bottom: 1px solid #eef0f2;
  }}
  td.num {{
    text-align: right; color: var(--muted); background: #f6f8fa;
    user-select: none; border-right: 1px solid var(--border);
  }}
  tr.delete td.cell-left {{ background: var(--red-bg); }}
  tr.delete td.num:first-child {{ background: var(--red-num); }}
  tr.insert td.cell-right {{ background: var(--green-bg); }}
  tr.insert td.num:last-child {{ background: var(--green-num); }}
  tr.replace td.cell-left {{ background: var(--red-bg); }}
  tr.replace td.cell-right {{ background: var(--green-bg); }}
  tr.replace td.num:first-child {{ background: var(--red-num); }}
  tr.replace td.num:last-child {{ background: var(--green-num); }}
  mark.w-del {{ background: var(--red-mark); text-decoration: line-through; padding: 0 1px; }}
  mark.w-ins {{ background: var(--green-mark); padding: 0 1px; }}
  tr.current {{ outline: 2px solid #0969da; outline-offset: -2px; }}
  tr.equal:hover td {{ background: #f6f8fa; }}
</style>
</head>
<body>
<header>
  <div>
    <h1>Perbandingan Dokumen</h1>
    <div class="doc-names">{doc_a_name} &nbsp;→&nbsp; {doc_b_name}</div>
  </div>
  <div class="summary">
    <span class="badge added">+{added} ditambahkan</span>
    <span class="badge removed">-{removed} dihapus</span>
    <span class="badge changed">~{changed} diubah</span>
    <div class="nav">
      <button id="nav-prev" title="Perubahan sebelumnya">&uarr; Sebelumnya</button>
      <span id="nav-counter">0 / 0</span>
      <button id="nav-next" title="Perubahan berikutnya">Berikutnya &darr;</button>
    </div>
  </div>
</header>
<table>
  <colgroup>
    <col class="num"><col class="cell"><col class="num"><col class="cell">
  </colgroup>
  <tbody>
{rows}
  </tbody>
</table>
<script>
  const changeRows = Array.from(document.querySelectorAll('tr.changed'));
  let idx = -1;
  const counter = document.getElementById('nav-counter');
  function updateCounter() {{
    counter.textContent = changeRows.length ? (idx + 1) + ' / ' + changeRows.length : '0 / 0';
  }}
  function goTo(i) {{
    if (!changeRows.length) return;
    idx = (i + changeRows.length) % changeRows.length;
    changeRows.forEach(r => r.classList.remove('current'));
    changeRows[idx].classList.add('current');
    changeRows[idx].scrollIntoView({{behavior: 'smooth', block: 'center'}});
    updateCounter();
  }}
  document.getElementById('nav-next').addEventListener('click', () => goTo(idx + 1));
  document.getElementById('nav-prev').addEventListener('click', () => goTo(idx - 1));
  updateCounter();
</script>
</body>
</html>
"""


def _row_html(row_type, no_a, left_html, no_b, right_html):
    row_class = row_type + (" changed" if row_type != "equal" else "")
    return (
        f'<tr class="{row_class}">'
        f'<td class="num">{no_a if no_a is not None else ""}</td>'
        f'<td class="cell cell-left">{left_html}</td>'
        f'<td class="num">{no_b if no_b is not None else ""}</td>'
        f'<td class="cell cell-right">{right_html}</td>'
        f"</tr>"
    )


def compare_documents(path_a, path_b, output_html_path):
    lines_a = extract_lines(path_a)
    lines_b = extract_lines(path_b)

    rows, counts = _build_rows(lines_a, lines_b)
    rows_html = "\n".join(_row_html(*row) for row in rows)

    html_doc = _HTML_TEMPLATE.format(
        doc_a_name=html.escape(os.path.basename(path_a)),
        doc_b_name=html.escape(os.path.basename(path_b)),
        added=counts["added"],
        removed=counts["removed"],
        changed=counts["changed"],
        rows=rows_html,
    )
    with open(output_html_path, "w", encoding="utf-8") as f:
        f.write(html_doc)

    return counts, output_html_path


def main():
    parser = argparse.ArgumentParser(
        description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument("--a", required=True, help="Path dokumen versi lama (.docx/.pdf)")
    parser.add_argument("--b", required=True, help="Path dokumen versi baru (.docx/.pdf)")
    parser.add_argument("--output", required=True, help="Path file HTML hasil perbandingan")
    args = parser.parse_args()

    counts, output_path = compare_documents(args.a, args.b, args.output)
    print(f"Dokumen A : {args.a}")
    print(f"Dokumen B : {args.b}")
    print(f"Ditambahkan : {counts['added']}")
    print(f"Dihapus     : {counts['removed']}")
    print(f"Diubah      : {counts['changed']}")
    print(f"Hasil       : {output_path}")


if __name__ == "__main__":
    main()
